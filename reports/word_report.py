"""
Generador de informes Word para procesamiento BRP.
"""

from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List, Optional

import pandas as pd
import matplotlib
matplotlib.use('Agg')  # Backend sin GUI
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reports.audit_log import AuditLog


class InformeWord:
    """
    Generador de informes Word para procesamiento BRP.

    Crea documentos con resumen ejecutivo, logs de auditoría,
    gráficos y comparaciones mensuales.
    """

    def __init__(self):
        self.doc = Document()
        self._configure_styles()

    def _configure_styles(self) -> None:
        """Configura estilos del documento."""
        # Título del documento
        style = self.doc.styles['Title']
        style.font.size = Pt(24)
        style.font.bold = True

        # Heading 1
        style = self.doc.styles['Heading 1']
        style.font.size = Pt(16)
        style.font.bold = True

        # Heading 2
        style = self.doc.styles['Heading 2']
        style.font.size = Pt(14)
        style.font.bold = True

    def generar(
        self,
        mes: str,
        df_resultado: pd.DataFrame,
        audit_log: AuditLog,
        comparacion: Optional[Dict[str, Any]] = None
    ) -> BytesIO:
        """
        Genera el informe Word completo.

        Args:
            mes: Identificador del mes procesado
            df_resultado: DataFrame con resultados del procesamiento
            audit_log: Log de auditoría del procesamiento
            comparacion: Resultado de comparación con mes anterior (opcional)

        Returns:
            Buffer con el documento Word
        """
        # 1. Portada
        self._agregar_portada(mes)

        # 2. Resumen ejecutivo
        self._agregar_resumen(df_resultado, mes)

        # 3. Distribución BRP
        self._agregar_seccion_distribucion(df_resultado)

        # 4. Gráficos
        self._agregar_graficos(df_resultado)

        # 5. Docentes EIB (posibles)
        self._agregar_seccion_eib(df_resultado, audit_log)

        # 6. Valores inusuales y advertencias
        self._agregar_valores_inusuales(audit_log)

        # 7. Log de procesamiento
        self._agregar_logs(audit_log)

        # 8. Comparación mensual (si aplica)
        if comparacion:
            self._agregar_comparacion(comparacion)

        return self._to_buffer()

    def _agregar_portada(self, mes: str) -> None:
        """Agrega portada del informe."""
        # Espacio superior
        for _ in range(3):
            self.doc.add_paragraph()

        # Título
        titulo = self.doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("INFORME DE DISTRIBUCIÓN BRP")
        run.font.size = Pt(28)
        run.font.bold = True

        # Subtítulo
        subtitulo = self.doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitulo.add_run("Bonificación de Reconocimiento Profesional")
        run.font.size = Pt(16)

        # Espacio
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Mes
        mes_p = self.doc.add_paragraph()
        mes_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = mes_p.add_run(f"Período: {mes}")
        run.font.size = Pt(18)
        run.font.bold = True

        # Fecha de generación
        fecha = self.doc.add_paragraph()
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fecha.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        run.font.size = Pt(12)

        # Salto de página
        self.doc.add_page_break()

    def _agregar_resumen(self, df: pd.DataFrame, mes: str) -> None:
        """Agrega resumen ejecutivo."""
        self.doc.add_heading("1. Resumen Ejecutivo", level=1)

        # Calcular métricas
        brp_sep = df['BRP_SEP'].sum() if 'BRP_SEP' in df.columns else 0
        brp_pie = df['BRP_PIE'].sum() if 'BRP_PIE' in df.columns else 0
        brp_normal = df['BRP_NORMAL'].sum() if 'BRP_NORMAL' in df.columns else 0
        brp_total = brp_sep + brp_pie + brp_normal

        # Identificar columna RUT
        rut_col = 'RUT_NORM' if 'RUT_NORM' in df.columns else None
        if not rut_col:
            for col in df.columns:
                if 'rut' in col.lower():
                    rut_col = col
                    break

        total_docentes = df[rut_col].nunique() if rut_col else len(df)

        # Identificar columna RBD
        rbd_col = None
        for col in df.columns:
            if 'rbd' in col.lower():
                rbd_col = col
                break
        total_rbds = df[rbd_col].nunique() if rbd_col else 0

        # Docentes con BRP = 0 (posibles EIB)
        docentes_eib = len(df[df['BRP_TOTAL'] == 0]) if 'BRP_TOTAL' in df.columns else 0

        # Tabla de resumen
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'

        data = [
            ('Concepto', 'Valor'),
            ('Período Procesado', mes),
            ('Total Docentes', f"{total_docentes:,}"),
            ('Total Establecimientos', f"{total_rbds:,}"),
            ('BRP Total Distribuido', f"${brp_total:,.0f}"),
            ('Docentes con BRP $0 (posibles EIB)', f"{docentes_eib:,}"),
            ('Fecha de Procesamiento', datetime.now().strftime('%d/%m/%Y'))
        ]

        for i, (concepto, valor) in enumerate(data):
            row = table.rows[i]
            row.cells[0].text = concepto
            row.cells[1].text = str(valor)

            # Encabezado en negrita
            if i == 0:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        self.doc.add_paragraph()

    def _agregar_seccion_distribucion(self, df: pd.DataFrame) -> None:
        """Agrega sección de distribución BRP."""
        self.doc.add_heading("2. Distribución por Tipo de Subvención", level=1)

        brp_sep = df['BRP_SEP'].sum() if 'BRP_SEP' in df.columns else 0
        brp_pie = df['BRP_PIE'].sum() if 'BRP_PIE' in df.columns else 0
        brp_normal = df['BRP_NORMAL'].sum() if 'BRP_NORMAL' in df.columns else 0
        brp_total = brp_sep + brp_pie + brp_normal

        # Tabla de distribución
        table = self.doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        headers = ['Subvención', 'Monto', 'Porcentaje']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        data = [
            ('SEP', brp_sep, brp_sep/brp_total*100 if brp_total > 0 else 0),
            ('PIE', brp_pie, brp_pie/brp_total*100 if brp_total > 0 else 0),
            ('NORMAL', brp_normal, brp_normal/brp_total*100 if brp_total > 0 else 0),
            ('TOTAL', brp_total, 100)
        ]

        for i, (tipo, monto, pct) in enumerate(data, 1):
            table.rows[i].cells[0].text = tipo
            table.rows[i].cells[1].text = f"${monto:,.0f}"
            table.rows[i].cells[2].text = f"{pct:.1f}%"

        self.doc.add_paragraph()

        # Desglose por concepto
        self.doc.add_heading("2.1 Desglose por Concepto", level=2)

        recon_cols = ['BRP_RECONOCIMIENTO_SEP', 'BRP_RECONOCIMIENTO_PIE', 'BRP_RECONOCIMIENTO_NORMAL']
        tramo_cols = ['BRP_TRAMO_SEP', 'BRP_TRAMO_PIE', 'BRP_TRAMO_NORMAL']

        recon_total = sum(df[col].sum() for col in recon_cols if col in df.columns)
        tramo_total = sum(df[col].sum() for col in tramo_cols if col in df.columns)

        p = self.doc.add_paragraph()
        p.add_run(f"Reconocimiento Profesional: ").bold = True
        p.add_run(f"${recon_total:,.0f}")

        p = self.doc.add_paragraph()
        p.add_run(f"Tramo: ").bold = True
        p.add_run(f"${tramo_total:,.0f}")

        self.doc.add_paragraph()

    def _agregar_graficos(self, df: pd.DataFrame) -> None:
        """Agrega gráficos al informe."""
        self.doc.add_heading("3. Visualización", level=1)

        # Gráfico de distribución por subvención
        brp_sep = df['BRP_SEP'].sum() if 'BRP_SEP' in df.columns else 0
        brp_pie = df['BRP_PIE'].sum() if 'BRP_PIE' in df.columns else 0
        brp_normal = df['BRP_NORMAL'].sum() if 'BRP_NORMAL' in df.columns else 0

        if brp_sep + brp_pie + brp_normal > 0:
            # Gráfico de torta
            fig, ax = plt.subplots(figsize=(6, 4))
            valores = [brp_sep, brp_pie, brp_normal]
            etiquetas = ['SEP', 'PIE', 'NORMAL']
            colores = ['#3b82f6', '#10b981', '#f59e0b']

            # Filtrar valores cero
            datos = [(v, e, c) for v, e, c in zip(valores, etiquetas, colores) if v > 0]
            if datos:
                valores_f, etiquetas_f, colores_f = zip(*datos)
                ax.pie(valores_f, labels=etiquetas_f, autopct='%1.1f%%',
                       colors=colores_f, startangle=90)
                ax.set_title('Distribución por Tipo de Subvención')

                # Guardar en buffer
                buf = BytesIO()
                plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                buf.seek(0)
                plt.close()

                self.doc.add_picture(buf, width=Inches(4.5))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

    def _agregar_seccion_eib(self, df: pd.DataFrame, audit_log: AuditLog) -> None:
        """Agrega sección de docentes EIB."""
        self.doc.add_heading("4. Docentes con BRP $0 (Posibles EIB)", level=1)

        # Obtener docentes con BRP = 0
        if 'BRP_TOTAL' in df.columns:
            df_eib = df[df['BRP_TOTAL'] == 0]
        else:
            df_eib = pd.DataFrame()

        if df_eib.empty:
            self.doc.add_paragraph("No se detectaron docentes con BRP $0.")
        else:
            self.doc.add_paragraph(
                f"Se identificaron {len(df_eib)} docentes con BRP igual a $0. "
                "Estos pueden corresponder a docentes del programa EIB "
                "(Educación Intercultural Bilingüe) u otras situaciones especiales."
            )

            # Mostrar primeros 20
            self.doc.add_paragraph()

            # Identificar columnas
            rut_col = 'RUT_NORM' if 'RUT_NORM' in df_eib.columns else None
            nombre_col = None
            rbd_col = None

            for col in df_eib.columns:
                if 'rut' in col.lower() and not rut_col:
                    rut_col = col
                if 'nombre' in col.lower():
                    nombre_col = col
                if 'rbd' in col.lower():
                    rbd_col = col

            if rut_col and len(df_eib) > 0:
                table = self.doc.add_table(rows=min(len(df_eib), 20) + 1, cols=3)
                table.style = 'Table Grid'

                # Encabezados
                table.rows[0].cells[0].text = 'RUT'
                table.rows[0].cells[1].text = 'Nombre'
                table.rows[0].cells[2].text = 'RBD'

                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Datos
                for i, (_, row) in enumerate(df_eib.head(20).iterrows(), 1):
                    table.rows[i].cells[0].text = str(row.get(rut_col, ''))
                    table.rows[i].cells[1].text = str(row.get(nombre_col, '')) if nombre_col else ''
                    table.rows[i].cells[2].text = str(row.get(rbd_col, '')) if rbd_col else ''

                if len(df_eib) > 20:
                    self.doc.add_paragraph(
                        f"(Mostrando 20 de {len(df_eib)} docentes)"
                    )

        self.doc.add_paragraph()

    def _agregar_valores_inusuales(self, audit_log: AuditLog) -> None:
        """Agrega sección de valores inusuales y advertencias."""
        self.doc.add_heading("5. Valores Inusuales y Advertencias", level=1)

        warnings = audit_log.get_warnings()
        errors = audit_log.get_errors()

        if not warnings and not errors:
            self.doc.add_paragraph("No se detectaron valores inusuales ni advertencias.")
            return

        if errors:
            self.doc.add_heading("5.1 Errores", level=2)
            for entry in errors[:10]:
                p = self.doc.add_paragraph()
                run = p.add_run(f"[ERROR] {entry.mensaje}")
                run.font.color.rgb = RGBColor(255, 0, 0)

        if warnings:
            self.doc.add_heading("5.2 Advertencias", level=2)
            for entry in warnings[:20]:
                p = self.doc.add_paragraph()
                run = p.add_run(f"[WARNING] {entry.mensaje}")
                run.font.color.rgb = RGBColor(255, 165, 0)

        self.doc.add_paragraph()

    def _agregar_logs(self, audit_log: AuditLog) -> None:
        """Agrega log de procesamiento."""
        self.doc.add_heading("6. Log de Procesamiento", level=1)

        summary = audit_log.get_summary()

        p = self.doc.add_paragraph()
        p.add_run(f"Total de eventos: ").bold = True
        p.add_run(f"{summary.get('total', 0)}")

        p = self.doc.add_paragraph()
        p.add_run(f"Errores: ").bold = True
        p.add_run(f"{summary.get('errores', 0)}")

        p = self.doc.add_paragraph()
        p.add_run(f"Advertencias: ").bold = True
        p.add_run(f"{summary.get('advertencias', 0)}")

        # Mostrar últimos eventos
        if audit_log.entries:
            self.doc.add_heading("6.1 Eventos del Procesamiento", level=2)

            table = self.doc.add_table(rows=min(len(audit_log.entries), 15) + 1, cols=3)
            table.style = 'Table Grid'

            headers = ['Hora', 'Nivel', 'Mensaje']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                for paragraph in table.rows[0].cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            for i, entry in enumerate(audit_log.entries[:15], 1):
                table.rows[i].cells[0].text = entry.timestamp.strftime('%H:%M:%S')
                table.rows[i].cells[1].text = entry.nivel
                table.rows[i].cells[2].text = entry.mensaje[:50] + ('...' if len(entry.mensaje) > 50 else '')

        self.doc.add_paragraph()

    def _agregar_comparacion(self, comparacion: Dict[str, Any]) -> None:
        """Agrega sección de comparación mensual."""
        self.doc.add_page_break()
        self.doc.add_heading("7. Comparación con Mes Anterior", level=1)

        resumen = comparacion.get('resumen', {})
        mes_ant = comparacion.get('mes_anterior', '')
        mes_act = comparacion.get('mes_actual', '')

        self.doc.add_paragraph(f"Comparación: {mes_ant} vs {mes_act}")

        # Tabla de resumen
        table = self.doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'

        headers = ['Concepto', mes_ant, mes_act, 'Cambio']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        data = [
            ('Total Docentes',
             str(resumen.get('docentes_anterior', 0)),
             str(resumen.get('docentes_actual', 0)),
             f"{resumen.get('docentes_actual', 0) - resumen.get('docentes_anterior', 0):+d}"),
            ('Docentes Nuevos', '-',
             str(resumen.get('docentes_nuevos', 0)), '+'),
            ('Docentes Salieron',
             str(resumen.get('docentes_salieron', 0)), '-', '-'),
            ('BRP Total',
             f"${resumen.get('brp_anterior', 0):,.0f}",
             f"${resumen.get('brp_actual', 0):,.0f}",
             f"{resumen.get('cambio_brp_pct', 0):+.1f}%"),
            ('BRP SEP',
             f"${resumen.get('sep_anterior', 0):,.0f}",
             f"${resumen.get('sep_actual', 0):,.0f}", ''),
            ('BRP PIE',
             f"${resumen.get('pie_anterior', 0):,.0f}",
             f"${resumen.get('pie_actual', 0):,.0f}", ''),
            ('BRP NORMAL',
             f"${resumen.get('normal_anterior', 0):,.0f}",
             f"${resumen.get('normal_actual', 0):,.0f}", '')
        ]

        for i, row_data in enumerate(data, 1):
            for j, val in enumerate(row_data):
                table.rows[i].cells[j].text = val

        self.doc.add_paragraph()

        # Docentes nuevos
        docentes_nuevos = comparacion.get('docentes_nuevos', [])
        if docentes_nuevos:
            self.doc.add_heading("7.1 Docentes Nuevos", level=2)
            self.doc.add_paragraph(f"Se incorporaron {len(docentes_nuevos)} docentes:")

            for doc in docentes_nuevos[:10]:
                self.doc.add_paragraph(
                    f"- {doc.get('nombre', '')} ({doc.get('rut', '')}): "
                    f"${doc.get('brp_total', 0):,.0f}",
                    style='List Bullet'
                )

            if len(docentes_nuevos) > 10:
                self.doc.add_paragraph(f"... y {len(docentes_nuevos) - 10} más")

        # Docentes que salieron
        docentes_salieron = comparacion.get('docentes_salieron', [])
        if docentes_salieron:
            self.doc.add_heading("7.2 Docentes que Salieron", level=2)
            self.doc.add_paragraph(f"Salieron {len(docentes_salieron)} docentes:")

            for doc in docentes_salieron[:10]:
                self.doc.add_paragraph(
                    f"- {doc.get('nombre', '')} ({doc.get('rut', '')})",
                    style='List Bullet'
                )

            if len(docentes_salieron) > 10:
                self.doc.add_paragraph(f"... y {len(docentes_salieron) - 10} más")

        # Cambios significativos de monto
        cambios_montos = comparacion.get('cambios_montos', [])
        if cambios_montos:
            self.doc.add_heading("7.3 Cambios Significativos de Monto (>10%)", level=2)

            for cambio in cambios_montos[:10]:
                self.doc.add_paragraph(
                    f"- {cambio.get('nombre', '')} ({cambio.get('rut', '')}): "
                    f"${cambio.get('monto_anterior', 0):,.0f} -> "
                    f"${cambio.get('monto_actual', 0):,.0f} "
                    f"({cambio.get('cambio_porcentaje', 0):+.1f}%)",
                    style='List Bullet'
                )

        self.doc.add_paragraph()

    def _to_buffer(self) -> BytesIO:
        """Convierte el documento a buffer."""
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
